from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import comtypes.client
import os

# Crear documento
doc = Document()
output_dir = r"D:\dev\portafolio"
output_docx = os.path.join(output_dir, "Hoja_de_vida_Oscar_Goez.docx")
output_pdf = os.path.join(output_dir, "public\Hoja_de_vida_Oscar_Goez.pdf")


# ---------- CONFIGURAR ESTILOS ----------
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ---------- ENCABEZADO ----------
title = doc.add_heading("Oscar Alberto Goez Henao", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph("Automatización • Python • RPA • Data Analysis • React")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

contact = doc.add_paragraph("Bello, Antioquia | 3105290842 | oscargoezhenao@gmail.com")
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# ---------- PERFIL PROFESIONAL ----------
doc.add_heading("Perfil Profesional", level=1)
doc.add_paragraph(
    "Tecnólogo en Análisis y Desarrollo de Sistemas de Información con formación en "
    "automatización de procesos, análisis de datos y desarrollo web. Combino experiencia "
    "en logística y operaciones con habilidades técnicas en Python, RPA y análisis de datos "
    "para crear soluciones que optimizan la eficiencia y la gestión empresarial."
)

# ---------- COMPETENCIAS BLANDAS ----------
doc.add_heading("Competencias Blandas", level=1)
doc.add_paragraph(
    "• Pensamiento analítico: Capacidad para descomponer procesos complejos y encontrar puntos de mejora. "
    "Esta habilidad me ha permitido automatizar tareas repetitivas y optimizar flujos de trabajo.\n"
    "• Orientación a la eficiencia: Mi experiencia en logística me enseñó a detectar desperdicios de tiempo y recursos, "
    "lo cual aplico en el desarrollo de automatizaciones que mejoran la productividad.\n"
    "• Resolución de problemas: Enfrento desafíos técnicos y operativos con pensamiento estructurado, "
    "proponiendo soluciones prácticas y sostenibles.\n"
    "• Trabajo en equipo: Acostumbrado a colaborar en entornos logísticos y tecnológicos, "
    "aportando comunicación y compromiso.\n"
    "• Adaptabilidad y aprendizaje continuo: La transición del entorno operativo al tecnológico "
    "refleja mi capacidad de aprendizaje y actualización constante en herramientas digitales."
)

# ---------- EXPERIENCIA LABORAL ----------
doc.add_heading("Experiencia Laboral", level=1)

# Freelance
doc.add_paragraph(
    "Freelance — Automatización de Procesos (2023 – Actualidad)\n"
    "• Diseño e implementación de bots en Python para tareas repetitivas y consultas web.\n"
    "• Clasificación automática de documentos PDF y reportes.\n"
    "• Arquitectura y desarrollo completo de sistema CRM para gestión de clientes, cotizaciones y facturación\n"
    "• Diseño y desarrollo completo de aplicación web SPA para gestión inteligente de inventario doméstico\n"
    "Consolida mi experiencia práctica en automatización, análisis de datos y desarrollo de soluciones funcionales."
)


# Altipal
doc.add_paragraph(
    "Operario de Montacargas — Altipal S.A.S (Septiembre 2023 – Marzo 2025)\n"
    "• Organización y logística del almacén.\n"
    "• Operación de montacargas asegurando el correcto flujo de la operación.\n"
    "• Control y cuidado del inventario.\n"
    "Esta experiencia reforzó mi capacidad de organización, precisión y mejora continua, "
    "habilidades que hoy aplico al desarrollo tecnológico."
)



# Enterdev
doc.add_paragraph(
    "Practicante Analista y Parametrizador RPA — Enterdev S.A.S (Julio 2022 – Enero 2023)\n"
    "• Análisis y parametrización de robots de automatización (RPA) según flujos BPMN definidos.\n"
    "• Participación en documentación técnica y validación de procesos automatizados.\n"
    "• Colaboración con el equipo de desarrollo en integración y pruebas de bots.\n"
    "Esta experiencia fortaleció mis bases en automatización de procesos empresariales con enfoque en eficiencia y calidad."
)

# ---------- EDUCACIÓN ----------
doc.add_heading("Educación", level=1)
doc.add_paragraph(
    "Servicio Nacional de Aprendizaje (SENA) — Tecnólogo en Análisis y Desarrollo de Sistemas de Información (2023)"
)
doc.add_paragraph("Cisco Networking Academy — Fundamentos de Python 1 y 2 (2025)")
doc.add_paragraph("MIT Professional Education — Internet de las Cosas (2025)")
doc.add_paragraph("IE University — Introducción a la Ciencia de Datos (2025)")
doc.add_paragraph("Curso de Excel Avanzado (2025)")
doc.add_paragraph("Marketing Automation (G-Talent, 2025)")
doc.add_paragraph("Microsoft Copilot (2025)")

# ---------- HABILIDADES TÉCNICAS ----------
doc.add_heading("Habilidades Técnicas", level=1)
doc.add_paragraph(
    "• Automatización con Python, Selenium\n"
    "• Manipulación y análisis de datos con Pandas y Excel\n"
    "• Desarrollo web con React y Node.js\n"
    "• Control de versiones con GitHub"
)

# ---------- PROYECTOS DESTACADOS ----------
doc.add_heading("Proyectos Destacados", level=1)
doc.add_paragraph(
    "• Bot de Automatización Excel → Web → Reporte: lectura de archivos Excel, consultas online y envío automático de reportes.\n"
    "• Clasificador de Facturas PDF: organización y clasificación automática de documentos.\n"
    "• Sitio web profesional desarrollado para la empresa HHO Ascensores & Malacates. https://ascensoreshho.web.app/\n"
    "• Sistema de Gestión HHO Ascensores: Plataforma web completa para administración de clientes, documentos financieros y métricas empresariales\n"
    "• Sistema de Inventario Inteligente para Hogares: Aplicación web completa para gestión automatizada de productos del hogar\n"
)

# ---------- PORTAFOLIO Y GITHUB ----------

section_heading = doc.add_heading("Portafolio y GitHub", level=1)
section_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph(
    "📂 Portafolio Web: https://oscargoezh.web.app\n"
    "💻 GitHub: https://github.com/oscargoez\n\n"
    "Estos enlaces permiten explorar proyectos prácticos de automatización, análisis de datos y desarrollo web "
    "realizados con Python, RPA y React."
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------- GUARDAR DOCX ----------
output_docx = os.path.join(output_dir, "Oscar_Goez_CV.docx")
doc.save(output_docx)
print(f"✅ Hoja de vida generada: {output_docx}")




# ---------- CONVERTIR A PDF ----------


if os.path.exists(output_docx):
    print("Convirtiendo a PDF...")
    def convert_docx_to_pdf(docx_path, pdf_path):
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()

    convert_docx_to_pdf(output_docx, output_pdf)
    print(f"✅ PDF generado correctamente: {output_pdf}")
else:
    print("❌ No se encontró el archivo .docx para convertir.")
